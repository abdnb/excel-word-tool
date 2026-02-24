#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Excel 批量生成 Word 文档 - 支持字段映射
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import pandas as pd
from docx import Document
import os
import sys
import threading
from datetime import datetime
import re


class ExcelToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Excel-Word工具")
        self.root.geometry("1100x750")
        self.root.resizable(True, True)
        
        # 设置窗口图标
        try:
            # 获取程序所在目录
            if getattr(sys, 'frozen', False):
                # 打包后的 exe
                application_path = sys._MEIPASS
            else:
                # 开发环境
                application_path = os.path.dirname(os.path.abspath(__file__))
            
            icon_path = os.path.join(application_path, 'icon.ico')
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
        except Exception as e:
            print(f"设置图标失败: {e}")
            pass  # 如果设置失败就使用默认图标
        
        # 设置主题色 - 简洁现代风格
        self.colors = {
            'primary': '#2C3E50',      # 深灰蓝 - 主色
            'secondary': '#34495E',    # 中灰蓝
            'accent': '#3498DB',       # 清爽蓝 - 强调色
            'success': '#27AE60',      # 绿色
            'bg': '#F5F6FA',           # 浅灰背景（更柔和）
            'white': '#FFFFFF',        # 白色
            'text': '#2C3E50',         # 文字颜色
            'text_light': '#7F8C8D',   # 浅色文字
            'border': '#E1E8ED'        # 边框颜色（更淡）
        }
        
        # 设置窗口背景色
        self.root.configure(bg=self.colors['bg'])
        
        # 配置样式
        self.setup_styles()
        
        # 变量
        self.excel_files = []
        self.template_file = ""
        self.output_dir = ""
        self.skip_rows = tk.IntVar(value=1)  # 改为默认跳过1行
        self.placeholders = []  # 模板中的占位符
        self.excel_columns = []  # Excel 列名
        self.field_mapping = {}  # 字段映射关系
        self.filter_conditions = []  # 过滤条件列表
        
        self.create_widgets()
    
    def setup_styles(self):
        """配置 ttk 样式"""
        style = ttk.Style()
        style.theme_use('clam')
        
        # 配置 Notebook（标签页）样式
        style.configure('TNotebook', background=self.colors['bg'], borderwidth=0)
        style.configure('TNotebook.Tab', 
                       background=self.colors['white'],
                       foreground=self.colors['text'],
                       padding=[20, 10],
                       font=('Microsoft YaHei UI', 10))
        style.map('TNotebook.Tab',
                 background=[('selected', self.colors['accent'])],
                 foreground=[('selected', 'white')])
        
        # 配置 Frame 样式
        style.configure('TFrame', background=self.colors['bg'])
        style.configure('Card.TFrame', background=self.colors['white'], relief='flat')
        
        # 配置 Label 样式
        style.configure('TLabel', 
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Microsoft YaHei UI', 9))
        style.configure('Title.TLabel',
                       font=('Microsoft YaHei UI', 12, 'bold'),
                       foreground=self.colors['primary'])
        style.configure('Card.TLabel',
                       background=self.colors['white'],
                       foreground=self.colors['text'],
                       font=('Microsoft YaHei UI', 9))
        
        # 配置 Button 样式
        style.configure('TButton',
                       font=('Microsoft YaHei UI', 9),
                       borderwidth=0,
                       relief='flat',
                       padding=[12, 6])
        style.configure('Accent.TButton',
                       background=self.colors['accent'],
                       foreground='white',
                       font=('Microsoft YaHei UI', 10, 'bold'),
                       padding=[20, 10])
        style.map('Accent.TButton',
                 background=[('active', '#2980B9')])
        
        # 配置 Entry 样式
        style.configure('TEntry',
                       fieldbackground='white',
                       borderwidth=1,
                       relief='solid')
        
        # 配置 LabelFrame 样式
        style.configure('TLabelframe',
                       background=self.colors['white'],
                       borderwidth=0,
                       relief='flat')
        style.configure('TLabelframe.Label',
                       background=self.colors['white'],
                       foreground=self.colors['primary'],
                       font=('Microsoft YaHei UI', 10, 'bold'))
        
        # 配置 Treeview 样式
        style.configure('Treeview',
                       background='white',
                       foreground=self.colors['text'],
                       fieldbackground='white',
                       font=('Microsoft YaHei UI', 9),
                       rowheight=32)
        style.configure('Treeview.Heading',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Microsoft YaHei UI', 9, 'bold'),
                       relief='flat',
                       borderwidth=0)
        style.map('Treeview.Heading',
                 background=[('active', self.colors['border'])])
        
        # 配置进度条样式
        style.configure('TProgressbar',
                       background=self.colors['accent'],
                       troughcolor=self.colors['border'],
                       borderwidth=0,
                       thickness=8)
        
    def create_widgets(self):
        # 创建主容器
        main_container = ttk.Frame(self.root, style='TFrame')
        main_container.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)
        
        # 内容区域（单页面）
        content_frame = ttk.Frame(main_container, style='TFrame', padding="25")
        content_frame.pack(fill=tk.BOTH, expand=True)
        content_frame.columnconfigure(0, weight=1)
        content_frame.columnconfigure(1, weight=1)
        content_frame.rowconfigure(0, weight=1)  # 顶部区域可扩展
        content_frame.rowconfigure(2, weight=1)  # 日志区域可扩展
        
        # 创建各个区域
        self.create_file_section_new(content_frame)
        self.create_mapping_section_new(content_frame)
        self.create_action_section_new(content_frame)
        self.create_log_section_new(content_frame)
    



    def create_file_section_new(self, parent):
        """创建文件选择区域（新版）"""
        # 创建卡片容器 - 左侧
        card_container = tk.Frame(parent, bg=self.colors['white'], relief='flat', bd=0)
        card_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 20), padx=(0, 10))
        card_container.columnconfigure(0, weight=1)
        card_container.rowconfigure(0, weight=1)
        
        # 添加阴影效果（通过边框模拟）
        shadow = tk.Frame(card_container, bg=self.colors['border'], bd=0)
        shadow.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        file_card = ttk.LabelFrame(shadow, text=" 文件选择 ", padding="20", style='TLabelframe')
        file_card.pack(fill=tk.BOTH, expand=True)
        file_card.columnconfigure(1, weight=1)
        
        # Excel 文件
        row = 0
        ttk.Label(file_card, text="Excel 文件", style='Card.TLabel').grid(row=row, column=0, sticky=tk.W, pady=8, padx=(0, 10))
        
        excel_frame = ttk.Frame(file_card, style='Card.TFrame')
        excel_frame.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=8)
        excel_frame.columnconfigure(0, weight=1)
        
        self.excel_listbox = tk.Listbox(excel_frame, height=3, 
                                        font=('Microsoft YaHei UI', 9),
                                        bg='white',
                                        relief='solid',
                                        borderwidth=1)
        self.excel_listbox.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        
        excel_btn_frame = ttk.Frame(excel_frame, style='Card.TFrame')
        excel_btn_frame.grid(row=0, column=1, sticky=(tk.N))
        
        ttk.Button(excel_btn_frame, text="添加", command=self.add_excel_files, width=10).pack(pady=2)
        ttk.Button(excel_btn_frame, text="删除选中", command=self.remove_selected_excel, width=10).pack(pady=2)
        ttk.Button(excel_btn_frame, text="清空全部", command=self.clear_excel_files, width=10).pack(pady=2)
        
        # 工作表选择
        row += 1
        ttk.Label(file_card, text="工作表", style='Card.TLabel').grid(row=row, column=0, sticky=tk.W, pady=8, padx=(0, 10))
        
        sheet_frame = ttk.Frame(file_card, style='Card.TFrame')
        sheet_frame.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=8)
        
        self.sheet_name = ttk.Combobox(sheet_frame, width=20, font=('Microsoft YaHei UI', 9), state='readonly')
        self.sheet_name.pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(sheet_frame, text="检测工作表", command=self.detect_sheets, width=12).pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Label(sheet_frame, text="（默认使用第一个工作表）", 
                 style='Card.TLabel',
                 foreground='#7F8C8D').pack(side=tk.LEFT)
        
        # Word 模板
        row += 1
        ttk.Label(file_card, text="Word 模板", style='Card.TLabel').grid(row=row, column=0, sticky=tk.W, pady=8, padx=(0, 10))
        
        template_frame = ttk.Frame(file_card, style='Card.TFrame')
        template_frame.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=8)
        template_frame.columnconfigure(0, weight=1)
        
        self.template_entry = ttk.Entry(template_frame, font=('Microsoft YaHei UI', 9))
        self.template_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        
        template_btn_frame = ttk.Frame(template_frame, style='Card.TFrame')
        template_btn_frame.grid(row=0, column=1)
        ttk.Button(template_btn_frame, text="浏览", command=self.select_template, width=10).pack(side=tk.LEFT, padx=2)
        ttk.Button(template_btn_frame, text="分析", command=self.analyze_template, width=10).pack(side=tk.LEFT, padx=2)
        
        # 输出目录
        row += 1
        ttk.Label(file_card, text="输出目录", style='Card.TLabel').grid(row=row, column=0, sticky=tk.W, pady=8, padx=(0, 10))
        
        output_frame = ttk.Frame(file_card, style='Card.TFrame')
        output_frame.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=8)
        output_frame.columnconfigure(0, weight=1)
        
        self.output_entry = ttk.Entry(output_frame, font=('Microsoft YaHei UI', 9))
        self.output_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        self.output_entry.insert(0, "output")
        ttk.Button(output_frame, text="浏览", command=self.select_output_dir, width=10).grid(row=0, column=1)
        
        # 跳过行数
        row += 1
        ttk.Label(file_card, text="跳过行数", style='Card.TLabel').grid(row=row, column=0, sticky=tk.W, pady=8, padx=(0, 10))
        
        skip_frame = ttk.Frame(file_card, style='Card.TFrame')
        skip_frame.grid(row=row, column=1, sticky=tk.W, pady=8)
        
        skip_entry = ttk.Entry(skip_frame, textvariable=self.skip_rows, width=10, font=('Microsoft YaHei UI', 9))
        skip_entry.pack(side=tk.LEFT, padx=(0, 10))
        ttk.Label(skip_frame, text="跳过前N行，下一行作为列名", 
                 style='Card.TLabel',
                 foreground='#7F8C8D').pack(side=tk.LEFT)
        ttk.Button(skip_frame, text="智能检测", command=self.auto_detect_skip_rows, width=10).pack(side=tk.LEFT, padx=(10, 0))
        ttk.Button(skip_frame, text="预览列名", command=self.preview_columns, width=10).pack(side=tk.LEFT, padx=(5, 0))
        
        # 文件名列
        row += 1
        ttk.Label(file_card, text="文件名列", style='Card.TLabel').grid(row=row, column=0, sticky=tk.W, pady=8, padx=(0, 10))
        
        filename_frame = ttk.Frame(file_card, style='Card.TFrame')
        filename_frame.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=8)
        
        self.filename_column = ttk.Combobox(filename_frame, width=20, font=('Microsoft YaHei UI', 9), state='readonly')
        self.filename_column.pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Label(filename_frame, text="（选择一列作为文件名，重复时自动加序号）", 
                 style='Card.TLabel',
                 foreground='#7F8C8D').pack(side=tk.LEFT)
        
        # 过滤条件
        row += 1
        ttk.Label(file_card, text="过滤条件", style='Card.TLabel').grid(row=row, column=0, sticky=(tk.W, tk.N), pady=8, padx=(0, 10))
        
        filter_frame = ttk.Frame(file_card, style='Card.TFrame')
        filter_frame.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=8)
        filter_frame.columnconfigure(0, weight=1)
        
        # 条件关系选择
        relation_frame = ttk.Frame(filter_frame, style='Card.TFrame')
        relation_frame.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Label(relation_frame, text="条件关系:", style='Card.TLabel').pack(side=tk.LEFT, padx=(0, 5))
        self.filter_relation = tk.StringVar(value="AND")
        ttk.Radiobutton(relation_frame, text="且（AND）- 同时满足所有条件", 
                       variable=self.filter_relation, value="AND").pack(side=tk.LEFT, padx=(0, 15))
        ttk.Radiobutton(relation_frame, text="或（OR）- 满足任一条件即可", 
                       variable=self.filter_relation, value="OR").pack(side=tk.LEFT)
        
        # 过滤条件列表框
        filter_list_frame = ttk.Frame(filter_frame, style='Card.TFrame')
        filter_list_frame.pack(fill=tk.BOTH, expand=True)
        
        self.filter_listbox = tk.Listbox(filter_list_frame, height=3, 
                                         font=('Microsoft YaHei UI', 9),
                                         bg='white',
                                         relief='solid',
                                         borderwidth=1)
        self.filter_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        filter_btn_frame = ttk.Frame(filter_list_frame, style='Card.TFrame')
        filter_btn_frame.pack(side=tk.LEFT)
        
        ttk.Button(filter_btn_frame, text="添加", command=self.add_filter_condition, width=10).pack(pady=2)
        ttk.Button(filter_btn_frame, text="删除选中", command=self.remove_filter_condition, width=10).pack(pady=2)
        ttk.Button(filter_btn_frame, text="清空", command=self.clear_filter_conditions, width=10).pack(pady=2)
        
        ttk.Label(filter_frame, text="（不添加条件则不过滤）", 
                 style='Card.TLabel',
                 foreground='#7F8C8D').pack(anchor=tk.W, pady=(5, 0))

    def create_mapping_section_new(self, parent):
        """创建字段映射区域（新版）"""
        # 创建卡片容器 - 右侧
        card_container = tk.Frame(parent, bg=self.colors['white'], relief='flat', bd=0)
        card_container.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 20), padx=(10, 0))
        card_container.columnconfigure(0, weight=1)
        card_container.rowconfigure(0, weight=1)
        
        # 添加阴影效果
        shadow = tk.Frame(card_container, bg=self.colors['border'], bd=0)
        shadow.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        mapping_card = ttk.LabelFrame(shadow, text=" 字段映射（双击编辑） ", padding="20", style='TLabelframe')
        mapping_card.pack(fill=tk.BOTH, expand=True)
        mapping_card.columnconfigure(0, weight=1)
        mapping_card.rowconfigure(0, weight=1)
        
        # 创建表格
        columns = ("placeholder", "excel_column", "preview")
        self.mapping_tree = ttk.Treeview(mapping_card, columns=columns, show="headings", height=6)
        
        self.mapping_tree.heading("placeholder", text="模板占位符")
        self.mapping_tree.heading("excel_column", text="Excel 列名")
        self.mapping_tree.heading("preview", text="预览值")
        
        self.mapping_tree.column("placeholder", width=280, anchor='center')
        self.mapping_tree.column("excel_column", width=280, anchor='center')
        self.mapping_tree.column("preview", width=220, anchor='center')
        
        # 滚动条
        scrollbar = ttk.Scrollbar(mapping_card, orient=tk.VERTICAL, command=self.mapping_tree.yview)
        self.mapping_tree.configure(yscrollcommand=scrollbar.set)
        
        self.mapping_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # 双击编辑
        self.mapping_tree.bind("<Double-1>", self.edit_mapping)
        
        # 按钮区域
        btn_frame = ttk.Frame(mapping_card, style='Card.TFrame')
        btn_frame.grid(row=1, column=0, columnspan=2, pady=(10, 0))
        
        ttk.Button(btn_frame, text="刷新映射表", command=self.refresh_mapping_table, width=12).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="自动匹配", command=self.auto_match_fields, width=12).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="清除映射", command=self.clear_mapping, width=12).pack(side=tk.LEFT, padx=5)

    def create_action_section_new(self, parent):
        """创建操作按钮区域（新版）"""
        # 创建卡片容器 - 横跨两列
        card_container = tk.Frame(parent, bg=self.colors['white'], relief='flat', bd=0)
        card_container.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 20))
        
        # 添加阴影效果
        shadow = tk.Frame(card_container, bg=self.colors['border'], bd=0)
        shadow.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        action_card = ttk.Frame(shadow, style='Card.TFrame', padding="20")
        action_card.pack(fill=tk.BOTH, expand=True)
        
        # 左侧：进度条
        progress_frame = ttk.Frame(action_card, style='Card.TFrame')
        progress_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 20))
        
        # 进度标签（包含百分比）
        progress_label_frame = ttk.Frame(progress_frame, style='Card.TFrame')
        progress_label_frame.pack(fill=tk.X, pady=(0, 5))
        
        ttk.Label(progress_label_frame, text="处理进度", style='Card.TLabel').pack(side=tk.LEFT)
        self.progress_label = ttk.Label(progress_label_frame, text="0%", style='Card.TLabel', foreground=self.colors['accent'])
        self.progress_label.pack(side=tk.RIGHT)
        
        self.progress = ttk.Progressbar(progress_frame, mode='determinate')
        self.progress.pack(fill=tk.X)
        
        # 右侧：按钮
        btn_container = ttk.Frame(action_card, style='Card.TFrame')
        btn_container.pack(side=tk.RIGHT)
        
        ttk.Button(btn_container, text="开始生成", 
                  command=self.start_generation, 
                  style='Accent.TButton',
                  width=15).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_container, text="打开输出目录", 
                  command=self.open_output_dir,
                  width=15).pack(side=tk.LEFT, padx=5)

    def create_log_section_new(self, parent):
        """创建日志区域（新版）"""
        # 创建卡片容器 - 横跨两列
        card_container = tk.Frame(parent, bg=self.colors['white'], relief='flat', bd=0)
        card_container.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S))
        card_container.columnconfigure(0, weight=1)
        card_container.rowconfigure(0, weight=1)
        
        # 添加阴影效果
        shadow = tk.Frame(card_container, bg=self.colors['border'], bd=0)
        shadow.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        log_card = ttk.LabelFrame(shadow, text=" 运行日志 ", padding="20", style='TLabelframe')
        log_card.pack(fill=tk.BOTH, expand=True)
        log_card.columnconfigure(0, weight=1)
        log_card.rowconfigure(0, weight=1)
        
        self.log_text = scrolledtext.ScrolledText(log_card, 
                                                 height=8,
                                                 font=('Consolas', 9),
                                                 bg='#FAFBFC',
                                                 fg=self.colors['text'],
                                                 relief='flat',
                                                 borderwidth=0,
                                                 wrap=tk.WORD)
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.log("欢迎使用 Excel-Word工具！")
        self.log("1. 添加 Excel 文件，点击'智能检测'自动设置跳过行数")
        self.log("2. 点击'预览列名'确认数据行数正确")
        self.log("3. 选择 Word 模板，点击'分析'查看占位符")
        self.log("4. 在字段映射表中双击配置字段对应关系")
        self.log("5. 点击'开始生成'批量生成文档")

    def log(self, message):
        """添加日志"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)
        self.root.update()
        
    def add_excel_files(self):
        """添加 Excel 文件（支持批量选择）"""
        files = filedialog.askopenfilenames(
            title="选择 Excel 文件（可多选）",
            filetypes=[("Excel 文件", "*.xlsx *.xls *.et"), ("所有文件", "*.*")]
        )
        added_count = 0
        for file in files:
            if file not in self.excel_files:
                self.excel_files.append(file)
                self.excel_listbox.insert(tk.END, os.path.basename(file))
                added_count += 1
        if added_count > 0:
            self.log(f"已添加 {added_count} 个 Excel 文件")
    
    def remove_selected_excel(self):
        """删除选中的 Excel 文件"""
        selection = self.excel_listbox.curselection()
        if not selection:
            messagebox.showwarning("提示", "请先选择要删除的文件")
            return
        
        # 从后往前删除，避免索引变化
        for index in reversed(selection):
            self.excel_listbox.delete(index)
            del self.excel_files[index]
        
        self.log(f"已删除 {len(selection)} 个文件")
    
    def clear_excel_files(self):
        """清空所有 Excel 文件"""
        if not self.excel_files:
            return
        
        result = messagebox.askyesno("确认", f"确定要清空全部 {len(self.excel_files)} 个文件吗？")
        if result:
            self.excel_files = []
            self.excel_listbox.delete(0, tk.END)
            self.log("已清空所有 Excel 文件")
        
    def detect_sheets(self):
        """检测 Excel 工作表"""
        if not self.excel_files:
            messagebox.showwarning("警告", "请先添加 Excel 文件")
            return
        
        try:
            excel_file = self.excel_files[0]
            sheet_names = None
            error_messages = []
            
            # 尝试多种引擎读取
            engines = ['openpyxl', 'xlrd', None]  # None 表示让 pandas 自动选择
            
            for engine in engines:
                try:
                    if engine:
                        xl_file = pd.ExcelFile(excel_file, engine=engine)
                    else:
                        xl_file = pd.ExcelFile(excel_file)
                    sheet_names = xl_file.sheet_names
                    break  # 成功读取，跳出循环
                except Exception as e:
                    error_messages.append(f"引擎 {engine or 'auto'}: {str(e)}")
                    continue
            
            if sheet_names is None:
                # 所有引擎都失败
                error_detail = "\n".join(error_messages)
                raise Exception(f"无法读取文件，请确保文件格式正确。\n\n尝试的引擎:\n{error_detail}\n\n建议：\n1. 在 Excel/WPS 中打开文件，另存为新的 .xlsx 文件\n2. 检查文件是否损坏\n3. 如果是 .et 文件，请转换为 .xlsx 格式")
            
            # 更新下拉框
            self.sheet_name['values'] = sheet_names
            if sheet_names:
                self.sheet_name.current(0)
            
            # 显示工作表列表
            sheets_text = "\n".join([f"{i+1}. {name}" for i, name in enumerate(sheet_names)])
            messagebox.showinfo("工作表列表", f"共 {len(sheet_names)} 个工作表:\n\n{sheets_text}")
            self.log(f"检测到 {len(sheet_names)} 个工作表: {', '.join(sheet_names)}")
            
        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.log(f"检测工作表失败: {e}")
    
    def preview_columns(self):
        """预览 Excel 列名"""
        if not self.excel_files:
            messagebox.showwarning("警告", "请先添加 Excel 文件")
            return
        
        try:
            excel_file = self.excel_files[0]
            df = self.read_excel_data(excel_file, self.skip_rows.get(), self.get_selected_sheet())
            self.excel_columns = df.columns.tolist()
            
            # 显示列名和数据行数
            columns_text = "\n".join([f"{i+1}. {col}" for i, col in enumerate(self.excel_columns)])
            messagebox.showinfo("Excel 列名", 
                f"共 {len(self.excel_columns)} 列:\n\n{columns_text}\n\n数据行数: {len(df)} 行")
            self.log(f"Excel 列名: {', '.join(self.excel_columns)}")
            self.log(f"数据行数: {len(df)} 行")
            
            # 更新文件名列下拉选项
            self.filename_column['values'] = ['（使用默认规则）'] + self.excel_columns
            if not self.filename_column.get():
                self.filename_column.current(0)
            
            # 刷新映射表
            self.refresh_mapping_table()
        except Exception as e:
            messagebox.showerror("错误", f"读取 Excel 失败:\n{e}\n\n请检查'跳过行数'设置是否正确")
            self.log(f"读取 Excel 失败: {e}")
    
    def auto_detect_skip_rows(self):
        """智能检测跳过行数"""
        if not self.excel_files:
            messagebox.showwarning("警告", "请先添加 Excel 文件")
            return
        
        try:
            excel_file = self.excel_files[0]
            sheet = self.get_selected_sheet()
            best_skip = 0
            best_score = 0
            
            # 尝试不同的跳过行数
            for skip in range(0, 10):
                try:
                    # 如果没有指定工作表，使用第一个
                    sheet_param = sheet if sheet else 0
                    
                    # 尝试多种引擎
                    df = None
                    for engine in ['openpyxl', 'xlrd', None]:
                        try:
                            if engine:
                                df = pd.read_excel(excel_file, sheet_name=sheet_param, header=skip, engine=engine)
                            else:
                                df = pd.read_excel(excel_file, sheet_name=sheet_param, header=skip)
                            break
                        except:
                            continue
                    
                    if df is None:
                        continue
                    
                    df = df.dropna(how='all')
                    
                    # 评分标准：
                    # 1. 列名不应该是数字
                    # 2. 列名不应该包含太多 Unnamed
                    # 3. 应该有数据行
                    score = 0
                    
                    # 检查列名
                    unnamed_count = sum(1 for col in df.columns if 'Unnamed' in str(col))
                    numeric_count = sum(1 for col in df.columns if isinstance(col, (int, float)))
                    
                    if unnamed_count < len(df.columns) * 0.5:  # Unnamed 列少于50%
                        score += 10
                    if numeric_count == 0:  # 没有数字列名
                        score += 10
                    if len(df) >= 2:  # 至少有2行数据
                        score += 10
                    
                    if score > best_score:
                        best_score = score
                        best_skip = skip
                except:
                    continue
            
            self.skip_rows.set(best_skip)
            self.log(f"智能检测建议：跳过 {best_skip} 行")
            
            # 自动预览
            self.preview_columns()
            
        except Exception as e:
            messagebox.showerror("错误", f"智能检测失败:\n{e}")
            self.log(f"智能检测失败: {e}")

    def select_template(self):
        """选择模板文件"""
        file = filedialog.askopenfilename(
            title="选择 Word 模板",
            filetypes=[("Word 文件", "*.docx"), ("所有文件", "*.*")]
        )
        if file:
            self.template_file = file
            self.template_entry.delete(0, tk.END)
            self.template_entry.insert(0, file)
            self.log(f"已选择模板: {os.path.basename(file)}")
            
    def analyze_template(self):
        """分析模板中的占位符"""
        template = self.template_entry.get()
        if not template or not os.path.exists(template):
            messagebox.showwarning("警告", "请先选择模板文件")
            return
        
        try:
            doc = Document(template)
            placeholders_set = set()
            
            # 从段落中提取占位符
            for paragraph in doc.paragraphs:
                placeholders_set.update(re.findall(r'\{\{(.+?)\}\}', paragraph.text))
            
            # 从表格中提取占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        placeholders_set.update(re.findall(r'\{\{(.+?)\}\}', cell.text))
            
            self.placeholders = sorted(list(placeholders_set))
            
            # 显示占位符
            if self.placeholders:
                placeholders_text = "\n".join([f"{i+1}. {{{{{p}}}}}" for i, p in enumerate(self.placeholders)])
                messagebox.showinfo("模板占位符", f"共 {len(self.placeholders)} 个占位符:\n\n{placeholders_text}")
                self.log(f"模板占位符: {', '.join(self.placeholders)}")
            else:
                messagebox.showinfo("模板占位符", "未找到占位符\n\n请在模板中使用 {{字段名}} 格式")
                self.log("模板中未找到占位符")
            
            # 刷新映射表
            self.refresh_mapping_table()
        except Exception as e:
            messagebox.showerror("错误", f"分析模板失败:\n{e}")
            self.log(f"分析模板失败: {e}")

    def select_output_dir(self):
        """选择输出目录"""
        directory = filedialog.askdirectory(title="选择输出目录")
        if directory:
            self.output_dir = directory
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, directory)
            self.log(f"输出目录: {directory}")
    
    def refresh_mapping_table(self):
        """刷新映射表"""
        # 清空表格
        for item in self.mapping_tree.get_children():
            self.mapping_tree.delete(item)
        
        if not self.placeholders:
            return
        
        # 获取预览数据
        preview_data = {}
        if self.excel_files:
            try:
                df = self.read_excel_data(self.excel_files[0], self.skip_rows.get(), self.get_selected_sheet())
                if len(df) > 0:
                    first_row = df.iloc[0]
                    for col in df.columns:
                        preview_data[col] = str(first_row[col])[:20]
            except:
                pass
        
        # 填充表格
        for placeholder in self.placeholders:
            excel_col = self.field_mapping.get(placeholder, "")
            preview = preview_data.get(excel_col, "") if excel_col else ""
            self.mapping_tree.insert("", tk.END, values=(placeholder, excel_col, preview))
        
        self.log("已刷新字段映射表")
    
    def edit_mapping(self, event):
        """编辑映射关系"""
        if not self.excel_columns:
            messagebox.showwarning("警告", "请先预览 Excel 列名")
            return
        
        # 获取选中的项
        selection = self.mapping_tree.selection()
        if not selection:
            return
        
        item = selection[0]
        values = self.mapping_tree.item(item, "values")
        placeholder = values[0]
        current_mapping = values[1]
        
        # 创建选择对话框
        dialog = tk.Toplevel(self.root)
        dialog.title(f"选择 Excel 列 - {placeholder}")
        dialog.geometry("400x500")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text=f"为占位符 {{{{{placeholder}}}}} 选择对应的 Excel 列:", 
                 padding="10").pack()
        
        # 列表框
        listbox = tk.Listbox(dialog, height=20)
        listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 添加选项
        listbox.insert(0, "（不映射）")
        for col in self.excel_columns:
            listbox.insert(tk.END, col)
        
        # 选中当前映射
        if current_mapping:
            try:
                idx = self.excel_columns.index(current_mapping) + 1
                listbox.selection_set(idx)
                listbox.see(idx)
            except:
                pass
        
        def on_select():
            selection = listbox.curselection()
            if selection:
                idx = selection[0]
                if idx == 0:
                    # 不映射
                    if placeholder in self.field_mapping:
                        del self.field_mapping[placeholder]
                else:
                    selected_col = self.excel_columns[idx - 1]
                    self.field_mapping[placeholder] = selected_col
                
                self.refresh_mapping_table()
                dialog.destroy()
        
        ttk.Button(dialog, text="确定", command=on_select).pack(pady=10)
        
        # 双击选择
        listbox.bind("<Double-1>", lambda e: on_select())

    def auto_match_fields(self):
        """自动匹配字段"""
        if not self.placeholders or not self.excel_columns:
            messagebox.showwarning("警告", "请先分析模板和预览 Excel 列名")
            return
        
        matched = 0
        for placeholder in self.placeholders:
            # 完全匹配
            if placeholder in self.excel_columns:
                self.field_mapping[placeholder] = placeholder
                matched += 1
            # 模糊匹配
            else:
                for col in self.excel_columns:
                    if placeholder.lower() in col.lower() or col.lower() in placeholder.lower():
                        self.field_mapping[placeholder] = col
                        matched += 1
                        break
        
        self.refresh_mapping_table()
        self.log(f"自动匹配完成，匹配了 {matched}/{len(self.placeholders)} 个字段")
        messagebox.showinfo("自动匹配", f"匹配了 {matched}/{len(self.placeholders)} 个字段")
    
    def clear_mapping(self):
        """清除映射"""
        self.field_mapping = {}
        self.refresh_mapping_table()
        self.log("已清除所有字段映射")
    
    def add_filter_condition(self):
        """添加过滤条件"""
        if not self.excel_columns:
            messagebox.showwarning("警告", "请先预览 Excel 列名")
            return
        
        # 创建对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("添加过滤条件")
        dialog.geometry("450x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 列名
        ttk.Label(dialog, text="列名:", padding="10").grid(row=0, column=0, sticky=tk.W)
        col_combo = ttk.Combobox(dialog, width=20, font=('Microsoft YaHei UI', 9), state='readonly')
        col_combo['values'] = self.excel_columns
        col_combo.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=10, pady=10)
        if self.excel_columns:
            col_combo.current(0)
        
        # 条件
        ttk.Label(dialog, text="条件:", padding="10").grid(row=1, column=0, sticky=tk.W)
        op_combo = ttk.Combobox(dialog, width=20, font=('Microsoft YaHei UI', 9), state='readonly')
        op_combo['values'] = ('等于', '不等于', '包含', '不包含', '为空', '不为空')
        op_combo.current(0)
        op_combo.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=10, pady=10)
        
        # 值
        ttk.Label(dialog, text="值:", padding="10").grid(row=2, column=0, sticky=tk.W)
        val_entry = ttk.Entry(dialog, width=22, font=('Microsoft YaHei UI', 9))
        val_entry.grid(row=2, column=1, sticky=(tk.W, tk.E), padx=10, pady=10)
        
        def on_confirm():
            col = col_combo.get()
            op = op_combo.get()
            val = val_entry.get().strip()
            
            if not col:
                messagebox.showwarning("警告", "请选择列名")
                return
            
            # 添加到列表
            condition = {'column': col, 'operator': op, 'value': val}
            self.filter_conditions.append(condition)
            
            # 显示在列表框
            display_text = f"{col} {op}"
            if op not in ('为空', '不为空'):
                display_text += f" '{val}'"
            self.filter_listbox.insert(tk.END, display_text)
            
            self.log(f"添加过滤条件: {display_text}")
            dialog.destroy()
        
        ttk.Button(dialog, text="确定", command=on_confirm).grid(row=3, column=0, columnspan=2, pady=20)
    
    def remove_filter_condition(self):
        """删除选中的过滤条件"""
        selection = self.filter_listbox.curselection()
        if not selection:
            messagebox.showwarning("提示", "请先选择要删除的条件")
            return
        
        for index in reversed(selection):
            self.filter_listbox.delete(index)
            del self.filter_conditions[index]
        
        self.log(f"已删除 {len(selection)} 个过滤条件")
    
    def clear_filter_conditions(self):
        """清空所有过滤条件"""
        if not self.filter_conditions:
            return
        
        result = messagebox.askyesno("确认", f"确定要清空全部 {len(self.filter_conditions)} 个过滤条件吗？")
        if result:
            self.filter_conditions = []
            self.filter_listbox.delete(0, tk.END)
            self.log("已清空所有过滤条件")
            
    def open_output_dir(self):
        """打开输出目录"""
        output_dir = self.output_entry.get()
        if os.path.exists(output_dir):
            os.startfile(output_dir)
        else:
            messagebox.showwarning("警告", "输出目录不存在")

    def start_generation(self):
        """开始生成"""
        if not self.excel_files:
            messagebox.showwarning("警告", "请先添加 Excel 文件")
            return
            
        template = self.template_entry.get()
        if not template or not os.path.exists(template):
            messagebox.showwarning("警告", "请选择有效的模板文件")
            return
            
        output_dir = self.output_entry.get()
        if not output_dir:
            messagebox.showwarning("警告", "请设置输出目录")
            return
        
        if not self.field_mapping:
            result = messagebox.askyesno("提示", "未配置字段映射，将使用完全匹配模式。\n\n是否继续？")
            if not result:
                return
            
        # 在新线程中执行生成任务
        thread = threading.Thread(target=self.generate_documents)
        thread.daemon = True
        thread.start()
        
    def generate_documents(self):
        """生成文档（在后台线程中执行）"""
        try:
            output_dir = self.output_entry.get()
            template = self.template_entry.get()
            skip_rows = self.skip_rows.get()
            
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # 计算总文档数
            total_docs = 0
            file_doc_counts = []
            for excel_file in self.excel_files:
                try:
                    df = self.read_excel_data(excel_file, skip_rows, self.get_selected_sheet())
                    df_filtered = self._apply_filter(df)
                    doc_count = len(df_filtered)
                    file_doc_counts.append(doc_count)
                    total_docs += doc_count
                except:
                    file_doc_counts.append(0)
            
            if total_docs == 0:
                messagebox.showwarning("提示", "没有符合条件的数据需要生成")
                return
            
            self.progress['maximum'] = total_docs
            self.progress['value'] = 0
            self.progress_label.config(text="0% (预计时间: 计算中...)")
            
            total_generated = 0
            start_time = datetime.now()
            
            for file_idx, excel_file in enumerate(self.excel_files):
                self.log(f"\n处理文件 ({file_idx + 1}/{len(self.excel_files)}): {os.path.basename(excel_file)}")
                
                try:
                    # 读取 Excel
                    df = self.read_excel_data(excel_file, skip_rows, self.get_selected_sheet())
                    self.log(f"  读取到 {len(df)} 条数据")
                    
                    # 应用过滤条件
                    df_filtered = self._apply_filter(df)
                    if len(df_filtered) < len(df):
                        self.log(f"  应用过滤条件后剩余 {len(df_filtered)} 条数据")
                    
                    # 为每个 Excel 创建子目录
                    excel_name = os.path.splitext(os.path.basename(excel_file))[0]
                    sub_output_dir = os.path.join(output_dir, excel_name)
                    if not os.path.exists(sub_output_dir):
                        os.makedirs(sub_output_dir)
                    
                    # 文件名计数器（用于处理重复文件名）
                    filename_counter = {}
                    
                    # 生成文档
                    for doc_idx, (index, row) in enumerate(df_filtered.iterrows()):
                        data_dict = {}
                        
                        # 使用字段映射
                        if self.field_mapping:
                            for placeholder, excel_col in self.field_mapping.items():
                                if excel_col in df.columns:
                                    value = row[excel_col]
                                    if pd.isna(value):
                                        value = ""
                                    else:
                                        value = self._format_value(value, excel_col)
                                    data_dict[placeholder] = value
                        else:
                            # 完全匹配模式
                            for col in df.columns:
                                value = row[col]
                                if pd.isna(value):
                                    value = ""
                                else:
                                    value = self._format_value(value, col)
                                data_dict[col] = value
                        
                        # 添加序号
                        data_dict['序号'] = index + 1
                        data_dict['index'] = index + 1
                        
                        # 生成文件名
                        filename_col = self.filename_column.get()
                        
                        if filename_col and filename_col != '（使用默认规则）':
                            # 使用指定列作为文件名
                            if filename_col in data_dict:
                                base_filename = str(data_dict[filename_col]).strip()
                            else:
                                base_filename = f"文档{total_generated + 1}"
                        else:
                            # 使用默认规则
                            dept_name = str(data_dict.get('委托部门名称', data_dict.get('部门', f'部门{index+1}'))).strip()
                            item_name = str(data_dict.get('受理事项名称', data_dict.get('事项', ''))).strip()
                            
                            if item_name:
                                base_filename = f"{dept_name}_{item_name}"
                            else:
                                base_filename = f"{dept_name}_{index+1}"
                        
                        # 清理文件名中的非法字符
                        for char in ['/', '\\', ':', '*', '?', '"', '<', '>', '|']:
                            base_filename = base_filename.replace(char, '_')
                        
                        # 处理重复文件名
                        if base_filename in filename_counter:
                            filename_counter[base_filename] += 1
                            output_filename = f"{base_filename}_{filename_counter[base_filename]}.docx"
                        else:
                            filename_counter[base_filename] = 0
                            output_filename = f"{base_filename}.docx"
                        
                        output_path = os.path.join(sub_output_dir, output_filename)
                        self.fill_template(template, output_path, data_dict)
                        total_generated += 1
                        
                        # 更新进度条和时间预估
                        self.progress['value'] = total_generated
                        progress_percent = int(total_generated / total_docs * 100)
                        
                        # 计算预计剩余时间
                        elapsed_time = (datetime.now() - start_time).total_seconds()
                        if total_generated > 0:
                            avg_time_per_doc = elapsed_time / total_generated
                            remaining_docs = total_docs - total_generated
                            estimated_remaining = int(avg_time_per_doc * remaining_docs)
                            
                            if estimated_remaining >= 60:
                                time_str = f"{estimated_remaining // 60}分{estimated_remaining % 60}秒"
                            else:
                                time_str = f"{estimated_remaining}秒"
                            
                            self.progress_label.config(text=f"{progress_percent}% (预计剩余: {time_str})")
                        else:
                            self.progress_label.config(text=f"{progress_percent}%")
                    
                    self.log(f"  完成！生成 {len(df_filtered)} 个文档")
                    
                except Exception as e:
                    self.log(f"  处理失败: {e}")
            
            self.progress_label.config(text="100% (已完成)")
            self.log(f"\n全部完成！共处理 {len(self.excel_files)} 个 Excel 文件，生成 {total_generated} 个 Word 文档")
            messagebox.showinfo("完成", f"批量生成完成！\n\n共生成 {total_generated} 个 Word 文档\n保存在: {output_dir}")
            
        except Exception as e:
            self.log(f"错误: {e}")
            messagebox.showerror("错误", f"生成失败:\n{e}")
    
    def _apply_filter(self, df):
        """应用过滤条件（支持 AND 和 OR 关系）"""
        # 如果没有设置过滤条件，返回原数据
        if not self.filter_conditions:
            return df
        
        relation = self.filter_relation.get()
        
        if relation == "AND":
            # AND 逻辑：所有条件都必须满足
            result_df = df.copy()
            
            for condition in self.filter_conditions:
                filter_col = condition['column']
                filter_op = condition['operator']
                filter_val = condition['value']
                
                # 检查列是否存在
                if filter_col not in result_df.columns:
                    self.log(f"  警告：过滤列 '{filter_col}' 不存在，跳过此条件")
                    continue
                
                try:
                    if filter_op == '等于':
                        result_df = result_df[result_df[filter_col].astype(str) == filter_val]
                    elif filter_op == '不等于':
                        result_df = result_df[result_df[filter_col].astype(str) != filter_val]
                    elif filter_op == '包含':
                        result_df = result_df[result_df[filter_col].astype(str).str.contains(filter_val, na=False)]
                    elif filter_op == '不包含':
                        result_df = result_df[~result_df[filter_col].astype(str).str.contains(filter_val, na=False)]
                    elif filter_op == '为空':
                        result_df = result_df[result_df[filter_col].isna() | (result_df[filter_col].astype(str).str.strip() == '')]
                    elif filter_op == '不为空':
                        result_df = result_df[result_df[filter_col].notna() & (result_df[filter_col].astype(str).str.strip() != '')]
                except Exception as e:
                    self.log(f"  警告：过滤条件 '{filter_col} {filter_op}' 应用失败 - {e}")
                    continue
            
            return result_df
        
        else:  # OR 逻辑
            # OR 逻辑：满足任一条件即可
            mask = pd.Series([False] * len(df), index=df.index)
            
            for condition in self.filter_conditions:
                filter_col = condition['column']
                filter_op = condition['operator']
                filter_val = condition['value']
                
                # 检查列是否存在
                if filter_col not in df.columns:
                    self.log(f"  警告：过滤列 '{filter_col}' 不存在，跳过此条件")
                    continue
                
                try:
                    if filter_op == '等于':
                        mask |= (df[filter_col].astype(str) == filter_val)
                    elif filter_op == '不等于':
                        mask |= (df[filter_col].astype(str) != filter_val)
                    elif filter_op == '包含':
                        mask |= df[filter_col].astype(str).str.contains(filter_val, na=False)
                    elif filter_op == '不包含':
                        mask |= ~df[filter_col].astype(str).str.contains(filter_val, na=False)
                    elif filter_op == '为空':
                        mask |= (df[filter_col].isna() | (df[filter_col].astype(str).str.strip() == ''))
                    elif filter_op == '不为空':
                        mask |= (df[filter_col].notna() & (df[filter_col].astype(str).str.strip() != ''))
                except Exception as e:
                    self.log(f"  警告：过滤条件 '{filter_col} {filter_op}' 应用失败 - {e}")
                    continue
            
            return df[mask]
    
    def _format_value(self, value, column_name=""):
        """格式化值，特别处理日期类型"""
        # 检查是否是日期时间类型
        if isinstance(value, pd.Timestamp) or hasattr(value, 'strftime'):
            # 如果列名是"今天"，格式化为 YYYY年M月D日
            if column_name == "今天":
                year = value.year
                month = value.month
                day = value.day
                return f'{year}年{month}月{day}日'
            else:
                # 其他日期列只显示日期部分，不显示时间
                return value.strftime('%Y-%m-%d')
        else:
            return value

    def read_excel_data(self, excel_file, skip_rows=3, sheet_name=None):
        """读取 Excel 数据（支持 .xlsx, .xls, .et 格式）"""
        # skip_rows 表示跳过前几行，然后下一行作为列名
        # 例如：skip_rows=1 表示跳过第1行，第2行是列名，第3行开始是数据
        
        # 如果没有指定工作表，使用第一个
        if sheet_name is None or sheet_name == '':
            sheet_name = 0
        
        df = None
        error_messages = []
        
        # 尝试多种引擎读取
        engines = ['openpyxl', 'xlrd', None]  # None 表示让 pandas 自动选择
        
        for engine in engines:
            try:
                if engine:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name, header=skip_rows, engine=engine)
                else:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name, header=skip_rows)
                break  # 成功读取，跳出循环
            except Exception as e:
                error_messages.append(f"引擎 {engine or 'auto'}: {str(e)}")
                continue
        
        if df is None:
            # 所有引擎都失败
            error_detail = "\n".join(error_messages)
            raise Exception(f"无法读取文件。\n\n尝试的引擎:\n{error_detail}\n\n建议：\n1. 在 Excel/WPS 中打开文件，另存为新的 .xlsx 文件\n2. 检查文件是否损坏")
        
        # 删除全空行
        df = df.dropna(how='all')
        
        # 删除列名为 NaN 或 Unnamed 的列
        df = df.loc[:, ~df.columns.astype(str).str.contains('^Unnamed')]
        
        return df
    
    def get_selected_sheet(self):
        """获取选中的工作表名称"""
        sheet = self.sheet_name.get()
        return sheet if sheet else None
        
    def fill_template(self, template_path, output_path, data_dict):
        """填充模板，统一应用字体格式"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        doc = Document(template_path)
        
        # 处理段落
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data_dict)
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data_dict)
        
        doc.save(output_path)
    
    def _replace_in_paragraph(self, paragraph, data_dict):
        """在段落中替换占位符，统一应用字体格式（支持跨run的占位符）"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        # 先检查段落文本中是否有占位符
        full_text = paragraph.text
        has_placeholder = False
        
        for key in data_dict.keys():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                has_placeholder = True
                break
        
        if not has_placeholder:
            return
        
        # 如果有占位符，需要处理跨run的情况
        # 方法：合并所有run的文本，替换后重新分配
        for key, value in data_dict.items():
            placeholder = f"{{{{{key}}}}}"
            value_str = str(value)
            
            if placeholder not in paragraph.text:
                continue
            
            # 保存第一个run的格式（用于新内容）
            if len(paragraph.runs) > 0:
                first_run = paragraph.runs[0]
            else:
                continue
            
            # 获取完整文本并替换
            full_text = paragraph.text
            new_text = full_text.replace(placeholder, value_str)
            
            # 如果文本没有变化，跳过
            if new_text == full_text:
                continue
            
            # 清空所有run
            for run in paragraph.runs:
                run.text = ""
            
            # 在第一个run中设置新文本
            first_run.text = new_text
            
            # 应用字体格式
            self._apply_default_font(first_run, new_text)
    
    def _apply_default_font(self, run, text):
        """应用默认字体：中文用方正仿宋_GBK，数字用Times New Roman，三号"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        # 三号字体 = 16磅
        font_size = Pt(16)
        
        # 检查文本内容
        has_chinese = any('\u4e00' <= char <= '\u9fff' for char in text)
        has_digit = any(char.isdigit() for char in text)
        
        if has_chinese and has_digit:
            # 混合内容：设置中文字体和西文字体
            run.font.name = '方正仿宋_GBK'
            run.font.size = font_size
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        elif has_digit:
            # 纯数字：使用 Times New Roman
            run.font.name = 'Times New Roman'
            run.font.size = font_size
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        else:
            # 纯中文或其他：使用方正仿宋_GBK
            run.font.name = '方正仿宋_GBK'
            run.font.size = font_size
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')


def main():
    root = tk.Tk()
    app = ExcelToWordApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
